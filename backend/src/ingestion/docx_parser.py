"""
docx_parser.py — Author: Suresh D R | AI Product Developer & Technology Mentor
Parses Word documents extracting paragraphs and tables separately.
"""
import re
from docx import Document

def parse_docx(file_path):
    doc = Document(file_path)
    paragraphs, tables = [], []
    current_section = "General"
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        is_heading = para.style and para.style.name.startswith("Heading")
        if is_heading:
            current_section = text
        paragraphs.append({
            "text": _clean(text),
            "style": para.style.name if para.style else "Normal",
            "section_name": current_section,
            "is_heading": bool(is_heading),
        })
    for idx, table in enumerate(doc.tables):
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        if rows:
            tables.append({
                "table_index": idx,
                "table_text": "\n".join([" | ".join(r) for r in rows]),
                "rows": rows, "row_count": len(rows), "col_count": len(rows[0]),
            })
    return paragraphs, tables

def _clean(text):
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()
